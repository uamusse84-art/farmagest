# -*- coding: utf-8 -*-
"""
Gera o relatorio final do projecto FarmaGest em formato Word (.docx).

    python tools\\gerar-relatorio.py

Requer: pip install python-docx
As imagens (diagramas UML e capturas de ecra) tem de existir em docs\\imagens.
"""

import os
import sys

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMAGENS = os.path.join(RAIZ, "docs", "imagens")
SAIDA = os.path.join(RAIZ, "docs", "Relatorio FarmaGest.docx")

VERDE = RGBColor(0x0F, 0x5C, 0x45)
CINZA = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()

# ---------------------------------------------------------------- estilos ---

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.4
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def definir_idioma(estilo, codigo="pt-PT"):
    rpr = estilo.element.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), codigo)


definir_idioma(normal)

for nivel, tamanho in ((1, 16), (2, 13), (3, 11.5)):
    estilo = doc.styles["Heading %d" % nivel]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(tamanho)
    estilo.font.bold = True
    estilo.font.color.rgb = VERDE
    estilo.paragraph_format.space_before = Pt(18 if nivel == 1 else 12)
    estilo.paragraph_format.space_after = Pt(6)
    estilo.paragraph_format.keep_with_next = True
    estilo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    definir_idioma(estilo)

for secao in doc.sections:
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.5)
    secao.left_margin = Cm(3.0)
    secao.right_margin = Cm(2.5)

LARGURA_UTIL = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

# ---------------------------------------------------------------- auxiliares ---

figura = {"n": 0}
tabela_n = {"n": 0}


def sombrear(celula, cor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), cor)
    celula._tc.get_or_add_tcPr().append(shd)


def p(texto="", **kw):
    par = doc.add_paragraph()
    if texto:
        run = par.add_run(texto)
        run.bold = kw.get("negrito", False)
        run.italic = kw.get("italico", False)
    if kw.get("centrado"):
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "espaco_antes" in kw:
        par.paragraph_format.space_before = Pt(kw["espaco_antes"])
    return par


def lista(itens, numerada=False):
    estilo = "List Number" if numerada else "List Bullet"
    for item in itens:
        par = doc.add_paragraph(style=estilo)
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.line_spacing = 1.25
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if isinstance(item, tuple):
            r = par.add_run(item[0])
            r.bold = True
            par.add_run(item[1])
        else:
            par.add_run(item)


def codigo(texto):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.6)
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(10)
    par.paragraph_format.line_spacing = 1.0
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    linhas = texto.strip("\n").split("\n")
    for i, linha in enumerate(linhas):
        run = par.add_run(linha)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.element.rPr.rFonts.set(qn("w:cs"), "Consolas")
        if i < len(linhas) - 1:
            run.add_break()
    pbdr = OxmlElement("w:pBdr")
    for lado in ("top", "left", "bottom", "right"):
        b = OxmlElement("w:%s" % lado)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), "CFD8DC")
        pbdr.append(b)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F5F7F8")
    ppr = par._p.get_or_add_pPr()
    ppr.append(pbdr)
    ppr.append(shd)
    return par


def imagem(ficheiro, legenda, largura=None, quebra_antes=False):
    caminho = os.path.join(IMAGENS, ficheiro)
    if not os.path.exists(caminho):
        print("  [AVISO] imagem em falta: %s" % ficheiro)
        return
    if quebra_antes:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(4)
    par.add_run().add_picture(caminho, width=largura or LARGURA_UTIL)
    figura["n"] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run("Figura %d - %s" % (figura["n"], legenda))
    r.font.size = Pt(9)
    r.font.color.rgb = CINZA
    r.italic = True


def tabela(cabecalho, linhas, legenda, larguras=None):
    tabela_n["n"] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run("Tabela %d - %s" % (tabela_n["n"], legenda))
    r.font.size = Pt(9)
    r.font.color.rgb = CINZA
    r.italic = True

    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, texto in enumerate(cabecalho):
        cel = t.rows[0].cells[i]
        cel.text = ""
        par = cel.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.space_after = Pt(2)
        run = par.add_run(texto)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(cel, "0F5C45")
    for j, linha in enumerate(linhas):
        celulas = t.add_row().cells
        for i, texto in enumerate(linha):
            celulas[i].text = ""
            par = celulas[i].paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.line_spacing = 1.1
            run = par.add_run(str(texto))
            run.font.size = Pt(9.5)
            if i == 0 and len(cabecalho) > 2:
                run.bold = True
            if j % 2 == 1:
                sombrear(celulas[i], "F2F6F4")
    if larguras:
        total = float(sum(larguras))
        for linha in t.rows:
            for i, cel in enumerate(linha.cells):
                cel.width = Cm(larguras[i] / total * 15.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def campo(par, instrucao, texto_provisorio=""):
    r1 = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r1._r.append(fc)
    r2 = par.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instrucao
    r2._r.append(it)
    r3 = par.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3._r.append(fc2)
    r4 = par.add_run(texto_provisorio)
    r5 = par.add_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r5._r.append(fc3)
    return r4


# ============================================================== CAPA =========

for _ in range(3):
    p()

t = p("FarmaGest", centrado=True)
r = t.runs[0]
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = VERDE

t = p("Sistema de Gestão de Farmácia", centrado=True)
t.runs[0].font.size = Pt(18)
t.runs[0].font.color.rgb = CINZA
t.paragraph_format.space_after = Pt(28)

t = p("Aplicação Web desenvolvida em PHP 8 sobre arquitectura MVC, "
      "com gestão de stock por lotes e abate automático segundo o critério FEFO",
      centrado=True, italico=True)
t.runs[0].font.size = Pt(11.5)
t.runs[0].font.color.rgb = CINZA
t.paragraph_format.space_after = Pt(60)

t = p("Relatório do Projecto", centrado=True, negrito=True)
t.runs[0].font.size = Pt(14)
t.paragraph_format.space_after = Pt(30)

for texto, negrito, tamanho in (
    ("Autor", False, 10),
    ("Nordino Elias Jossias Uamusse", True, 13),
    ("Código de estudante: 31240558", False, 11),
):
    par = p(texto, centrado=True, negrito=negrito)
    par.runs[0].font.size = Pt(tamanho)
    par.paragraph_format.space_after = Pt(4)

p().paragraph_format.space_after = Pt(24)

for texto in (
    "UniSCED - Universidade Aberta ISCED",
    "Cadeira: Desenvolvimento de Aplicativos Web Empresariais",
    "Moçambique, 2026",
):
    par = p(texto, centrado=True)
    par.runs[0].font.size = Pt(11)
    par.runs[0].font.color.rgb = CINZA
    par.paragraph_format.space_after = Pt(4)

# =============================================== SECÇÃO 2: corpo numerado ===

s2 = doc.add_section(WD_SECTION.NEW_PAGE)
s2.top_margin = Cm(2.5)
s2.bottom_margin = Cm(2.5)
s2.left_margin = Cm(3.0)
s2.right_margin = Cm(2.5)

# numeração a começar em 1 nesta secção
pg = OxmlElement("w:pgNumType")
pg.set(qn("w:start"), "1")
s2._sectPr.append(pg)

s2.footer.is_linked_to_previous = False
rodape = s2.footer.paragraphs[0]
rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
rodape.paragraph_format.space_after = Pt(0)
r = rodape.add_run("FarmaGest - Nordino Elias Jossias Uamusse (31240558)     |     ")
r.font.size = Pt(8.5)
r.font.color.rgb = CINZA
campo(rodape, " PAGE ", "1").font.size = Pt(8.5)
for run in rodape.runs:
    run.font.size = Pt(8.5)
    run.font.color.rgb = CINZA

doc.sections[0].footer.is_linked_to_previous = False  # capa sem rodapé

# ------------------------------------------------------------------ índice --

titulo = doc.add_paragraph()
titulo.paragraph_format.space_after = Pt(12)
r = titulo.add_run("Índice")
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = VERDE

par = doc.add_paragraph()
campo(par, 'TOC \\o "1-2" \\h \\z \\u',
      "Índice automático: no Word, clique com o botão direito e escolha "
      "«Actualizar campo» (ou prima F9) para o gerar.")

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================ 1. INTRODUÇÃO ==

doc.add_heading("1. Introdução", level=1)

doc.add_heading("1.1. Contextualização", level=2)
p("A farmácia comunitária é, em Moçambique, um dos pontos de contacto mais frequentes "
  "entre a população e o sistema de saúde. A gestão diária de uma farmácia envolve o "
  "controlo de um catálogo extenso de medicamentos, a recepção de mercadoria proveniente "
  "de vários fornecedores, o acompanhamento dos prazos de validade e o atendimento ao "
  "balcão, tudo isto sujeito a regras de rastreabilidade impostas pela natureza sanitária "
  "do produto transaccionado.")
p("Uma parte significativa dos estabelecimentos de menor dimensão continua a apoiar-se em "
  "registos manuais ou em folhas de cálculo isoladas. Este modelo é frágil: não impede a "
  "venda de produtos expirados, não garante que o stock apresentado corresponde ao stock "
  "real e não deixa um rasto de auditoria que permita apurar responsabilidades quando "
  "surge uma divergência de inventário.")
p("O presente trabalho descreve a concepção e a implementação do FarmaGest, uma aplicação "
  "Web de gestão de farmácia construída de raiz em PHP 8.2 sobre uma arquitectura "
  "Model-View-Controller (MVC), com base de dados relacional MySQL/MariaDB e interface "
  "responsiva assente em Bootstrap 5.")

doc.add_heading("1.2. Definição do problema", level=2)
p("O problema que motivou o projecto pode ser enunciado da seguinte forma: como assegurar "
  "que uma farmácia de pequena ou média dimensão consegue, simultaneamente, (i) vender "
  "sempre o lote que expira primeiro, (ii) impedir por construção a venda de produto "
  "expirado ou de stock inexistente, e (iii) manter um registo íntegro e auditável de "
  "todos os movimentos de existências?")
p("A dificuldade não está apenas em registar vendas. Está em garantir que o registo da "
  "venda e o abate de stock que dela decorre são indissociáveis: ou acontecem ambos, ou "
  "não acontece nenhum. Um sistema que grave a venda mas falhe o abate produz stock "
  "fantasma; um sistema que abata o stock mas falhe a venda produz perdas invisíveis.")

doc.add_heading("1.3. Objectivos", level=2)
p("Objectivo geral", negrito=True)
p("Desenvolver uma aplicação Web funcional para a gestão integrada de uma farmácia, "
  "cobrindo o catálogo de medicamentos, o controlo de existências por lote, o registo de "
  "vendas com abate automático de stock e a produção de relatórios de apoio à decisão.")
p("Objectivos específicos", negrito=True)
lista([
    "Modelar um esquema relacional normalizado que represente medicamentos, lotes, "
    "fornecedores, clientes, vendas e movimentos de stock, com integridade garantida ao "
    "nível da própria base de dados;",
    "Implementar uma arquitectura MVC própria, sem recurso a frameworks, de modo a "
    "demonstrar domínio dos mecanismos subjacentes (encaminhamento, camada de acesso a "
    "dados, renderização de vistas);",
    "Implementar o abate de stock segundo o critério FEFO (First Expired, First Out) "
    "dentro de uma transacção atómica com bloqueio de linhas;",
    "Garantir o controlo de acesso baseado em perfis (RBAC) com três níveis de "
    "permissão distintos;",
    "Aplicar as práticas correntes de segurança em aplicações Web: consultas "
    "parametrizadas, protecção contra CSRF, escape de saída, cifra de palavras-passe e "
    "limitação de tentativas de autenticação;",
    "Produzir a documentação técnica do sistema, incluindo os diagramas UML de casos de "
    "uso, de classes e de sequência;",
    "Validar o comportamento do sistema através de uma bateria de testes automatizados.",
])

doc.add_heading("1.4. Justificação", level=2)
p("A escolha do tema justifica-se por duas ordens de razões. Do ponto de vista prático, a "
  "gestão de validades é um problema real, com impacto sanitário e financeiro directo: um "
  "lote que expira em prateleira é perda total, e um lote expirado que é vendido é um "
  "risco para o doente. Do ponto de vista académico, o domínio farmacêutico obriga a "
  "lidar com concorrência, transacções, integridade referencial e auditoria, matérias "
  "centrais no desenvolvimento de aplicações Web empresariais.")

doc.add_heading("1.5. Metodologia", level=2)
p("O desenvolvimento seguiu um ciclo incremental organizado em sete etapas, cada uma "
  "concluída e verificada antes do início da seguinte:")
lista([
    "Estruturação do projecto e desenho da base de dados;",
    "Implementação do núcleo da aplicação (encaminhador, modelo base, sessão, "
    "autenticação, validação, tratamento de erros);",
    "Implementação dos modelos de domínio e dos controladores CRUD;",
    "Construção da interface responsiva;",
    "Testes da aplicação em execução;",
    "Produção dos diagramas UML;",
    "Redacção do relatório e preparação da entrega.",
], numerada=True)
p("A verificação foi feita a dois níveis: uma bateria de testes automatizados sobre a "
  "lógica de negócio, executada directamente contra a base de dados, e um teste de "
  "percurso que exercita todas as rotas da aplicação por HTTP com sessão autenticada.")

doc.add_heading("1.6. Estrutura do relatório", level=2)
p("Após esta introdução, o capítulo 2 apresenta a fundamentação teórica e as tecnologias "
  "adoptadas. O capítulo 3 formaliza os requisitos e os actores do sistema. O capítulo 4 "
  "descreve a modelação, incluindo os diagramas UML e o modelo de dados. O capítulo 5 "
  "detalha a implementação, com destaque para a regra de negócio central. O capítulo 6 "
  "apresenta os testes realizados e demonstra o sistema em funcionamento. O capítulo 7 "
  "conclui e identifica linhas de trabalho futuro.")

# ================================================= 2. FUNDAMENTAÇÃO TEÓRICA ==

doc.add_heading("2. Fundamentação teórica e tecnologias", level=1)

doc.add_heading("2.1. Gestão de existências em farmácia: o critério FEFO", level=2)
p("Na gestão de armazéns convencional, o critério mais divulgado para o consumo de "
  "existências é o FIFO (First In, First Out): a mercadoria que entrou primeiro é a "
  "primeira a sair. Este critério é adequado para bens cuja deterioração é proporcional "
  "ao tempo de permanência em armazém.")
p("O medicamento, porém, não obedece a essa lógica. Dois lotes do mesmo medicamento "
  "recebidos com uma semana de intervalo podem ter prazos de validade separados por anos, "
  "consoante a data de fabrico. Aplicar FIFO neste contexto pode levar a que um lote "
  "recebido mais tarde, mas com validade mais curta, permaneça em prateleira até expirar.")
p("O critério correcto é, por isso, o FEFO (First Expired, First Out): o lote a consumir "
  "é sempre aquele cuja data de validade está mais próxima, independentemente da data de "
  "entrada. É este o critério implementado no FarmaGest, e é ele que determina que o "
  "stock não possa residir na tabela de medicamentos - tem obrigatoriamente de residir na "
  "tabela de lotes, pois só o lote possui data de validade.")

doc.add_heading("2.2. A arquitectura Model-View-Controller", level=2)
p("O padrão MVC separa a aplicação em três responsabilidades. O Modelo encapsula os dados "
  "e as regras de negócio; a Vista é responsável exclusivamente pela apresentação; o "
  "Controlador recebe o pedido, coordena o modelo e escolhe a vista a devolver.")
p("No FarmaGest optou-se por implementar o padrão de raiz, sem recurso a uma framework. "
  "Esta decisão foi deliberada: o objectivo pedagógico do trabalho é demonstrar "
  "compreensão dos mecanismos que uma framework habitualmente esconde - o encaminhamento "
  "de pedidos, o carregamento automático de classes, a construção de consultas "
  "parametrizadas e o controlo do ciclo de vida da sessão. O fluxo de um pedido é o "
  "seguinte:")
codigo(
    "Pedido HTTP\n"
    "   -> public/index.php          ponto de entrada unico\n"
    "   -> app/bootstrap.php         autoloader, configuracao, sessao, erros\n"
    "   -> App\\Core\\Router          resolve metodo + caminho -> controlador\n"
    "      -> verifica autenticacao e perfil exigido pela rota\n"
    "   -> App\\Controllers\\...      valida entrada, invoca o modelo\n"
    "      -> App\\Models\\...        acesso a dados via PDO, regras de negocio\n"
    "   -> App\\Views\\...            renderiza HTML dentro de um layout\n"
    "   -> Resposta HTTP"
)

doc.add_heading("2.3. Tecnologias utilizadas", level=2)
tabela(
    ["Camada", "Tecnologia", "Justificação"],
    [
        ["Linguagem", "PHP 8.2",
         "Tipos estritos, propriedades promovidas no construtor, enumerações e "
         "readonly; ampla disponibilidade em alojamento partilhado."],
        ["Base de dados", "MySQL / MariaDB 10.4",
         "Motor InnoDB com suporte de transacções, chaves estrangeiras e bloqueio "
         "ao nível da linha - requisitos indispensáveis ao abate FEFO."],
        ["Acesso a dados", "PDO",
         "Camada de abstracção nativa com consultas preparadas reais "
         "(emulação desactivada), o que elimina a injecção de SQL."],
        ["Interface", "Bootstrap 5.3",
         "Sistema de grelha responsivo e componentes acessíveis; distribuído "
         "localmente para permitir funcionamento sem ligação à Internet."],
        ["Gráficos", "Chart.js 4",
         "Representação da evolução das vendas no painel de controlo."],
        ["Ícones", "Bootstrap Icons",
         "Conjunto de ícones vectoriais coerente com o Bootstrap."],
        ["Servidor", "Apache (XAMPP)",
         "Ambiente de desenvolvimento local; a aplicação funciona igualmente com "
         "o servidor embutido do PHP."],
    ],
    "Tecnologias adoptadas e respectiva justificação",
    larguras=[2.2, 2.6, 8.0],
)

# ================================================= 3. ANÁLISE DE REQUISITOS ==

doc.add_heading("3. Análise de requisitos", level=1)

doc.add_heading("3.1. Actores do sistema", level=2)
p("Foram identificados três actores, correspondentes aos três perfis de utilizador "
  "suportados pela aplicação:")
lista([
    ("Administrador - ", "acesso total ao sistema, incluindo a criação e a desactivação "
     "de contas de utilizador. Corresponde tipicamente ao proprietário ou ao director "
     "técnico do estabelecimento."),
    ("Farmacêutico - ", "gere o catálogo de medicamentos, as categorias, os "
     "fornecedores e as existências; regista entradas de mercadoria, efectua vendas, "
     "anula vendas e consulta relatórios."),
    ("Operador de caixa - ", "regista vendas e gere a carteira de clientes; consulta o "
     "catálogo e o stock em modo de leitura, sem poder alterá-los."),
])

doc.add_heading("3.2. Requisitos funcionais", level=2)
tabela(
    ["Código", "Requisito", "Perfis"],
    [
        ["RF01", "Autenticar o utilizador por e-mail e palavra-passe, terminando a sessão a pedido.", "Todos"],
        ["RF02", "Apresentar um painel de controlo com indicadores do dia, gráfico de vendas dos últimos sete dias e alertas de stock.", "Todos"],
        ["RF03", "Gerir o catálogo de medicamentos (criar, consultar, alterar e desactivar), com pesquisa e filtro por categoria.", "Administrador, Farmacêutico"],
        ["RF04", "Gerir categorias de medicamentos, impedindo a eliminação de categorias em uso.", "Administrador, Farmacêutico"],
        ["RF05", "Gerir fornecedores, com registo de NUIT e contactos.", "Administrador, Farmacêutico"],
        ["RF06", "Registar lotes de entrada com número de lote, quantidade, preço de compra, data de validade e fornecedor.", "Administrador, Farmacêutico"],
        ["RF07", "Efectuar ajustes manuais de stock, com registo obrigatório do motivo.", "Administrador, Farmacêutico"],
        ["RF08", "Sinalizar lotes expirados e lotes a expirar nos próximos noventa dias.", "Todos"],
        ["RF09", "Registar vendas com múltiplos itens, aplicando o abate de stock por FEFO.", "Todos"],
        ["RF10", "Impedir a venda de medicamentos inactivos, expirados ou com stock insuficiente.", "Sistema"],
        ["RF11", "Emitir um recibo imprimível para cada venda concluída.", "Todos"],
        ["RF12", "Anular uma venda, repondo integralmente o stock e conservando o registo original.", "Administrador, Farmacêutico"],
        ["RF13", "Gerir a carteira de clientes e consultar o respectivo histórico de compras.", "Todos"],
        ["RF14", "Produzir relatórios de vendas por período, com repartição por forma de pagamento e por operador.", "Todos"],
        ["RF15", "Produzir relatórios de existências e de movimentos de stock.", "Todos"],
        ["RF16", "Gerir contas de utilizador e respectivos perfis.", "Administrador"],
    ],
    "Requisitos funcionais do sistema",
    larguras=[1.3, 9.7, 3.5],
)

doc.add_heading("3.3. Requisitos não funcionais", level=2)
tabela(
    ["Código", "Requisito"],
    [
        ["RNF01", "Segurança: todas as consultas SQL têm de ser parametrizadas; as palavras-passe têm de ser cifradas com bcrypt; todos os pedidos que alteram estado têm de ser protegidos por um token anti-CSRF."],
        ["RNF02", "Integridade: o registo de uma venda e o abate de stock correspondente têm de ser atómicos - ou ambos são confirmados, ou nenhum o é."],
        ["RNF03", "Concorrência: duas vendas simultâneas do mesmo medicamento não podem consumir o mesmo lote duas vezes."],
        ["RNF04", "Auditoria: nenhuma venda é apagada; a anulação altera o estado e conserva o registo. Todo o movimento de stock fica registado com autor, data e motivo."],
        ["RNF05", "Usabilidade: a interface tem de ser utilizável em ecrãs de computador, tablet e telemóvel, e estar integralmente em português."],
        ["RNF06", "Portabilidade: a aplicação tem de funcionar em qualquer servidor com PHP 8.1 ou superior e MySQL 5.7 ou superior, sem dependências externas obrigatórias."],
        ["RNF07", "Manutenibilidade: separação estrita entre camadas; o código de apresentação não acede directamente à base de dados."],
    ],
    "Requisitos não funcionais do sistema",
    larguras=[1.3, 13.2],
)

doc.add_heading("3.4. Matriz de permissões", level=2)
tabela(
    ["Módulo", "Administrador", "Farmacêutico", "Caixa"],
    [
        ["Painel de controlo", "Total", "Total", "Total"],
        ["Medicamentos e categorias", "Total", "Total", "Leitura"],
        ["Fornecedores", "Total", "Total", "Leitura"],
        ["Lotes e ajustes de stock", "Total", "Total", "Leitura"],
        ["Registo de vendas", "Total", "Total", "Total"],
        ["Anulação de vendas", "Total", "Total", "Não"],
        ["Clientes", "Total", "Total", "Total"],
        ["Relatórios", "Total", "Total", "Total"],
        ["Utilizadores", "Total", "Não", "Não"],
    ],
    "Matriz de permissões por perfil",
    larguras=[5.5, 3.0, 3.0, 3.0],
)

doc.add_heading("3.5. Diagrama de casos de uso", level=2)
p("O diagrama seguinte sistematiza as interacções entre os três actores e as "
  "funcionalidades do sistema. As relações «include» assinalam comportamentos "
  "obrigatoriamente accionados pelo caso de uso principal - designadamente o abate de "
  "stock por FEFO e o registo do movimento correspondente, que são desencadeados sempre "
  "que uma venda é registada.")
imagem("01-casos-de-uso.png", "Diagrama de casos de uso do FarmaGest",
       largura=Cm(12.5))

# ==================================================== 4. MODELAÇÃO E DESENHO =

doc.add_heading("4. Modelação e desenho do sistema", level=1)

doc.add_heading("4.1. Modelo de dados", level=2)
p("A base de dados é composta por nove tabelas em motor InnoDB, com codificação "
  "utf8mb4, e por duas vistas de apoio aos relatórios. A decisão de desenho mais "
  "importante é a de o stock não residir na tabela de medicamentos: a quantidade "
  "disponível é sempre a soma das quantidades dos lotes válidos, porque só o lote possui "
  "data de validade e é essa data que determina a ordem de consumo.")
tabela(
    ["Tabela", "Finalidade", "Relações principais"],
    [
        ["utilizadores", "Contas de acesso, perfil e estado.", "Origem de vendas e movimentos"],
        ["categorias", "Classificação terapêutica dos medicamentos.", "1:N com medicamentos"],
        ["fornecedores", "Entidades que abastecem a farmácia.", "1:N com lotes"],
        ["medicamentos", "Catálogo: nome, dosagem, forma, preço de venda, stock mínimo.", "N:1 com categorias"],
        ["lotes", "Existências reais: número de lote, quantidade actual, validade.", "N:1 com medicamentos e fornecedores"],
        ["clientes", "Carteira de clientes, opcional na venda.", "1:N com vendas"],
        ["vendas", "Cabeçalho da venda: número, totais, pagamento, estado.", "N:1 com clientes e utilizadores"],
        ["itens_venda", "Linhas da venda, uma por lote consumido.", "N:1 com vendas, lotes e medicamentos"],
        ["movimentos_stock", "Registo de auditoria de todas as variações de existências.", "N:1 com lotes e utilizadores"],
    ],
    "Tabelas da base de dados e respectiva finalidade",
    larguras=[3.2, 6.8, 4.5],
)
p("A integridade é imposta na própria base de dados e não apenas na aplicação. Para além "
  "das chaves estrangeiras, foram declaradas restrições de verificação que tornam "
  "impossível gravar estados inválidos:")
codigo(
    "CONSTRAINT ck_medicamentos_preco  CHECK (preco_venda >= 0)\n"
    "CONSTRAINT ck_lotes_quantidade    CHECK (quantidade_atual >= 0)\n"
    "CONSTRAINT ck_vendas_desconto     CHECK (desconto >= 0 AND desconto <= subtotal)\n"
    "CONSTRAINT ck_itens_quantidade    CHECK (quantidade > 0)"
)
p("Note-se em particular a restrição sobre a quantidade dos lotes: mesmo que um defeito "
  "na aplicação tentasse abater mais unidades do que as existentes, a base de dados "
  "rejeitaria a operação e a transacção seria revertida na totalidade.")

doc.add_heading("4.2. Diagrama de classes", level=2)
p("O diagrama de classes representa a organização das camadas da aplicação. À esquerda "
  "situa-se o núcleo reutilizável, com destaque para a classe abstracta Model, da qual "
  "todos os modelos de domínio herdam as operações genéricas de leitura, escrita e "
  "paginação. À direita situam-se os modelos de domínio, com as respectivas "
  "especializações. As classes Venda e Lote concentram a lógica de negócio própria do "
  "sistema.")

A4_L = Cm(21.0)
A4_A = Cm(29.7)


def secao_paisagem(ficheiro):
    """Abre uma pagina em paisagem e devolve a largura a dar a imagem, ja
    limitada de modo a que a figura e a respectiva legenda caibam na pagina."""
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = A4_A, A4_L
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

    largura_util = s.page_width - s.left_margin - s.right_margin
    altura_util = s.page_height - s.top_margin - s.bottom_margin - Cm(1.4)  # legenda

    from PIL import Image  # noqa: WPS433  (dependencia de python-docx)
    with Image.open(os.path.join(IMAGENS, ficheiro)) as img:
        proporcao = img.width / img.height

    return min(largura_util, int(altura_util * proporcao))


def secao_retrato():
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = A4_L, A4_A
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(2.5)
    return s


largura_paisagem = secao_paisagem("02-diagrama-de-classes.png")
imagem("02-diagrama-de-classes.png",
       "Diagrama de classes: núcleo da aplicação e modelos de domínio",
       largura=largura_paisagem)

secao_retrato()
doc.add_heading("4.3. Diagrama de sequência: registo de uma venda", level=2)
p("O diagrama seguinte detalha o caso de uso mais complexo do sistema. Merecem destaque "
  "três aspectos: a transacção que envolve toda a operação, o ciclo exterior que percorre "
  "os itens do carrinho e o ciclo interior que percorre os lotes por ordem de validade "
  "até satisfazer a quantidade pedida. Qualquer excepção lançada no interior da "
  "transacção provoca a reversão integral.")

largura_paisagem = secao_paisagem("03-sequencia-registar-venda.png")
imagem("03-sequencia-registar-venda.png",
       "Diagrama de sequência do registo de uma venda com abate FEFO",
       largura=largura_paisagem)

secao_retrato()

# ================================================== 5. IMPLEMENTAÇÃO ========

doc.add_heading("5. Implementação", level=1)

doc.add_heading("5.1. Organização do código", level=2)
p("O código está organizado por responsabilidade. A pasta public é a única exposta ao "
  "servidor Web; todo o restante código reside fora da raiz pública, ficando inacessível "
  "por HTTP directo.")
codigo(
    "farmagest/\n"
    "  app/\n"
    "    Core/          nucleo: Router, Model, Database, Auth, Session,\n"
    "                   Csrf, Validator, Request, ErrorHandler, helpers\n"
    "    Models/        Medicamento, Lote, Venda, MovimentoStock,\n"
    "                   Categoria, Fornecedor, Cliente, Utilizador\n"
    "    Controllers/   um controlador por modulo funcional\n"
    "    Views/         vistas organizadas por modulo + layouts\n"
    "    bootstrap.php  arranque da aplicacao\n"
    "  config/          config.php (definicoes) e routes.php (rotas)\n"
    "  database/        schema.sql, seed.sql, seed_vendas.php\n"
    "  public/          index.php, .htaccess, assets/ (css, js, vendor)\n"
    "  storage/logs/    registo de erros\n"
    "  tests/           executar.php (unitarios) e crawl.ps1 (rotas)\n"
    "  tools/           geracao de diagramas, capturas e relatorio\n"
    "  docs/            diagramas, capturas de ecra e relatorio final"
)
tabela(
    ["Área", "Ficheiros", "Linhas"],
    [
        ["Núcleo (Core)", "16", "1 209"],
        ["Modelos de domínio", "8", "830"],
        ["Controladores", "10", "1 041"],
        ["Vistas", "30", "3 233"],
        ["Base de dados (SQL e povoamento)", "3", "424"],
        ["Configuração e rotas", "2", "130"],
        ["CSS e JavaScript próprios", "2", "337"],
        ["Testes automatizados", "2", "350"],
        ["Ferramentas de apoio", "3", "≈ 270"],
        ["Total (excluindo bibliotecas de terceiros)", "76", "≈ 7 800"],
    ],
    "Dimensão do código produzido",
    larguras=[8.5, 3.0, 3.0],
)

doc.add_heading("5.2. O encaminhador e o controlo de acesso", level=2)
p("As rotas são declaradas num único ficheiro, em que cada entrada associa um método HTTP "
  "e um caminho a um par controlador/acção, indicando ainda os perfis autorizados. O "
  "encaminhador converte os segmentos dinâmicos em parâmetros e verifica a autorização "
  "antes de instanciar o controlador, o que centraliza o controlo de acesso num único "
  "ponto:")
codigo(
    "$gestao = ['perfis' => ['farmaceutico']];\n"
    "$admin  = ['perfis' => ['administrador']];\n"
    "\n"
    "$router->get ('/medicamentos',            [MedicamentoController::class, 'indice']);\n"
    "$router->post('/medicamentos',            [MedicamentoController::class, 'guardar'], $gestao);\n"
    "$router->post('/vendas/{id}/anular',      [VendaController::class, 'anular'],        $gestao);\n"
    "$router->get ('/utilizadores',            [UtilizadorController::class, 'indice'],   $admin);"
)
p("A verificação de perfil é feita pelo método temPerfil da classe Auth. O perfil de "
  "administrador é aceite implicitamente em qualquer verificação, evitando ter de o "
  "enumerar em todas as rotas:")
codigo(
    "public static function temPerfil(string ...$perfis): bool\n"
    "{\n"
    "    $perfil = self::perfil();\n"
    "    return $perfil !== null\n"
    "        && ($perfil === 'administrador' || in_array($perfil, $perfis, true));\n"
    "}"
)

doc.add_heading("5.3. A camada de acesso a dados", level=2)
p("A classe abstracta Model concentra as operações comuns a todos os modelos. A ligação é "
  "estabelecida em PDO com a emulação de consultas preparadas desactivada, o que assegura "
  "que a preparação é feita pelo servidor de base de dados e que os valores nunca são "
  "interpolados na cadeia SQL:")
codigo(
    "PDO::ATTR_ERRMODE            => PDO::ERRMODE_EXCEPTION,\n"
    "PDO::ATTR_DEFAULT_FETCH_MODE => PDO::FETCH_ASSOC,\n"
    "PDO::ATTR_EMULATE_PREPARES   => false,"
)
p("Dois mecanismos merecem referência. O primeiro é a protecção contra atribuição em "
  "massa: cada modelo declara a lista de colunas preenchíveis, e qualquer chave submetida "
  "que não conste dessa lista é descartada em silêncio antes da escrita. O segundo diz "
  "respeito à ordenação: como os identificadores SQL não podem ser parametrizados por "
  "PDO, a cláusula de ordenação é validada contra uma expressão regular restritiva antes "
  "de ser concatenada, sendo substituída pela ordenação por omissão em caso de "
  "não conformidade.")

doc.add_heading("5.4. A regra de negócio central: o abate FEFO", level=2)
p("O método registar da classe Venda abre uma transacção, cria o cabeçalho da venda, "
  "delega o abate de cada item e só confirma no final. Qualquer excepção lançada em "
  "qualquer ponto do processo provoca a reversão de tudo o que já tinha sido escrito:")
codigo(
    "$bd->beginTransaction();\n"
    "try {\n"
    "    $numero  = $this->gerarNumero();\n"
    "    $vendaId = $this->criar([...]);\n"
    "\n"
    "    $subtotal = 0.0;\n"
    "    foreach ($itens as $item) {\n"
    "        $subtotal += $this->abaterStock($vendaId, $numero, $item, $utilizadorId);\n"
    "    }\n"
    "    // validacao do desconto e actualizacao dos totais ...\n"
    "    $bd->commit();\n"
    "    return $vendaId;\n"
    "} catch (Throwable $e) {\n"
    "    $bd->rollBack();\n"
    "    throw $e;\n"
    "}"
)
p("O coração do sistema é a selecção dos lotes. A consulta filtra os lotes com quantidade "
  "positiva e validade não expirada, ordena-os por data de validade crescente - é aqui "
  "que o critério FEFO se materializa - e bloqueia as linhas seleccionadas até ao fim da "
  "transacção:")
codigo(
    "SELECT id, numero_lote, quantidade_atual\n"
    "  FROM lotes\n"
    " WHERE medicamento_id = :id\n"
    "   AND quantidade_atual > 0\n"
    "   AND data_validade >= CURDATE()\n"
    " ORDER BY data_validade ASC, id ASC\n"
    "   FOR UPDATE"
)
p("A cláusula FOR UPDATE é indispensável e responde directamente ao requisito RNF03. Sem "
  "ela, duas vendas simultâneas do mesmo medicamento leriam ambas o mesmo lote com a mesma "
  "quantidade disponível e ambas concluiriam, com sucesso aparente, que havia stock "
  "suficiente - resultando em stock negativo. Com ela, a segunda transacção fica em "
  "espera até que a primeira confirme ou reverta, lendo então o valor já actualizado.")
p("Apurada a lista de lotes, a soma das quantidades é comparada com a quantidade pedida. "
  "Se for insuficiente, é lançada uma excepção de regra de negócio com uma mensagem "
  "explícita, e nada é gravado. Caso contrário, percorrem-se os lotes por ordem, "
  "retirando de cada um o mínimo entre o que falta atribuir e o que esse lote tem "
  "disponível. Uma consequência importante deste desenho é que a venda de uma única "
  "quantidade pode originar várias linhas em itens_venda, uma por cada lote consumido - "
  "o que preserva a rastreabilidade completa: para qualquer venda é sempre possível "
  "determinar exactamente de que lotes saiu cada unidade.")
p("Cada lote consumido gera igualmente um movimento de stock do tipo «saída», com "
  "quantidade negativa, referência ao número da venda e identificação do operador, "
  "satisfazendo o requisito de auditoria RNF04.")

doc.add_heading("5.5. Anulação de vendas", level=2)
p("A anulação obedece ao mesmo princípio de atomicidade. A venda não é eliminada: o seu "
  "estado passa a «anulada», é registado o motivo, e o stock é reposto lote a lote, "
  "gerando movimentos de entrada correspondentes. Uma venda já anulada não pode ser "
  "anulada de novo, e as vendas anuladas deixam de contar para a receita apresentada nos "
  "relatórios, mantendo-se contudo visíveis para efeitos de auditoria.")

doc.add_heading("5.6. Segurança", level=2)
p("As medidas de segurança implementadas foram as seguintes:")
lista([
    ("Injecção de SQL - ", "todas as consultas usam parâmetros ligados, com a emulação "
     "de preparação desactivada. Os identificadores que não podem ser parametrizados "
     "(nomes de tabela e de coluna na validação, cláusulas de ordenação) são validados "
     "contra listas ou expressões regulares restritivas."),
    ("Cross-Site Scripting - ", "toda a saída é escapada por uma função auxiliar que "
     "aplica htmlspecialchars com a codificação e as marcas explícitas."),
    ("Cross-Site Request Forgery - ", "é gerado um token por sessão, incluído em todos "
     "os formulários e verificado com comparação em tempo constante em todos os pedidos "
     "que não sejam GET."),
    ("Palavras-passe - ", "cifradas com bcrypt através de password_hash, e verificadas "
     "com password_verify. A palavra-passe em claro nunca é armazenada nem registada."),
    ("Força bruta - ", "após cinco tentativas falhadas, a autenticação fica bloqueada "
     "durante cinco minutos, com indicação do tempo em falta."),
    ("Fixação de sessão - ", "o identificador de sessão é regenerado no momento da "
     "autenticação; os cookies são marcados HttpOnly e SameSite=Lax."),
    ("Sessões órfãs - ", "a cada pedido é reconfirmado que a conta continua activa; se "
     "tiver sido desactivada entretanto, a sessão é terminada de imediato."),
    ("Exposição de erros - ", "em ambiente de produção o utilizador vê uma página de "
     "erro genérica, sendo o detalhe encaminhado exclusivamente para o ficheiro de "
     "registo."),
])

doc.add_heading("5.7. Interface do utilizador", level=2)
p("A interface foi construída sobre Bootstrap 5, com uma folha de estilo própria que "
  "define a identidade visual verde do sistema. O esquema é composto por uma barra "
  "superior com a identificação do utilizador e uma barra lateral de navegação que "
  "colapsa em ecrãs estreitos. As mensagens de sucesso e de erro circulam através da "
  "sessão, sendo apresentadas uma única vez após o redireccionamento.")
p("O ecrã de registo de venda é o mais elaborado do sistema. O carrinho é gerido no "
  "navegador por JavaScript, que apresenta o stock disponível de cada medicamento e "
  "calcula os totais em tempo real; contudo, todos os valores são recalculados no "
  "servidor a partir dos preços em base de dados, pelo que nenhum cálculo feito no "
  "cliente é aceite como verdadeiro.")

# ======================================================= 6. TESTES ==========

doc.add_heading("6. Testes e demonstração", level=1)

doc.add_heading("6.1. Testes da lógica de negócio", level=2)
p("Foi desenvolvido um pequeno arnês de testes que executa asserções directamente contra "
  "a base de dados, verificando o comportamento efectivo do sistema e não apenas o das "
  "funções isoladas. O conjunto é auto-limpante: as vendas e os movimentos criados "
  "durante os testes são removidos no final, e o estado inicial da base de dados é "
  "reposto e verificado.")
tabela(
    ["Grupo de testes", "Aspectos verificados", "Resultado"],
    [
        ["Configuração e ligação", "Leitura do ambiente, ligação PDO, existência das tabelas e das vistas.", "Passou"],
        ["Autenticação e perfis", "Credenciais válidas e inválidas, conta inactiva, matriz de permissões, bloqueio por tentativas.", "Passou"],
        ["Validação de dados", "Regras de obrigatoriedade, formatos, unicidade, datas futuras e passadas, força da palavra-passe.", "Passou"],
        ["Catálogo e existências", "Criação de medicamentos e lotes, cálculo do stock, detecção de lotes expirados e a expirar.", "Passou"],
        ["Venda com abate FEFO", "Ordem de consumo dos lotes, consumo através de vários lotes, recusa por stock insuficiente e por medicamento inactivo, cálculo dos totais.", "Passou"],
        ["Anulação de venda", "Mudança de estado, registo do motivo, reposição integral do stock, impossibilidade de anular duas vezes, exclusão da receita.", "Passou"],
        ["Integridade referencial", "Categoria em uso, medicamento com lotes, número de lote duplicado.", "Passou"],
    ],
    "Grupos de teste e respectivos resultados",
    larguras=[3.6, 8.4, 2.5],
)
p("O resultado da execução foi de 48 asserções executadas, 48 passadas e nenhuma "
  "falhada.")
codigo(
    "== Venda: abate de stock por FEFO ==\n"
    "  [OK]    o lote consumido foi o de validade mais proxima\n"
    "  [OK]    a venda atravessa varios lotes quando necessario\n"
    "  [OK]    venda com stock insuficiente e recusada\n"
    "  [OK]    medicamento inativo nao pode ser vendido\n"
    "  [OK]    o stock foi abatido exactamente na quantidade vendida\n"
    "  [OK]    foi registado um movimento de saida por cada lote consumido\n"
    "\n"
    "------------------------------------------------------------\n"
    "Total: 48 | Passados: 48 | Falhados: 0"
)

doc.add_heading("6.2. Teste de percurso das rotas", level=2)
p("Complementarmente, foi construído um teste que percorre por HTTP as cinquenta rotas de "
  "leitura da aplicação, com sessão previamente autenticada, e verifica não apenas o "
  "código de estado devolvido mas também a ausência, no corpo da resposta, de qualquer "
  "vestígio de erro do PHP. O resultado foi de cinquenta rotas testadas sem qualquer "
  "falha.")
codigo(
    "  [OK]    /relatorios/vendas\n"
    "  [OK]    /relatorios/stock\n"
    "  [OK]    /relatorios/movimentos\n"
    "\n"
    "Rotas testadas: 50 | Falhas: 0"
)

doc.add_heading("6.3. Demonstração do sistema", level=2)
p("As figuras seguintes documentam o sistema em funcionamento, com a base de dados "
  "povoada com dados de demonstração representativos: catálogo de medicamentos "
  "distribuídos por categorias terapêuticas, lotes com validades escalonadas e um "
  "histórico de vendas.")

demos = [
    ("ecras/01-login.png",
     "Ecrã de autenticação, com indicação das contas de demonstração"),
    ("ecras/02-painel.png",
     "Painel de controlo: indicadores do dia, evolução das vendas e alertas de stock"),
    ("ecras/03-medicamentos.png",
     "Catálogo de medicamentos, com pesquisa, filtro por categoria e indicação do stock"),
    ("ecras/04-lotes-validade.png",
     "Lotes a expirar nos próximos noventa dias, ordenados por proximidade da validade"),
    ("ecras/05-registar-venda.png",
     "Ecrã de registo de venda, com carrinho e cálculo de totais em tempo real"),
    ("ecras/06-recibo.png",
     "Recibo de venda em formato imprimível"),
    ("ecras/07-relatorio-vendas.png",
     "Relatório de vendas por período, com repartição por forma de pagamento e operador"),
    ("ecras/08-utilizadores.png",
     "Gestão de utilizadores, acessível apenas ao perfil de administrador"),
]
for ficheiro, legenda in demos:
    largura = Cm(10.0) if "recibo" in ficheiro else None
    imagem(ficheiro, legenda, largura=largura)

# ==================================================== 7. CONCLUSÕES =========

doc.add_heading("7. Conclusões e trabalho futuro", level=1)

doc.add_heading("7.1. Conclusões", level=2)
p("O objectivo geral proposto foi alcançado: o FarmaGest é uma aplicação Web funcional e "
  "completa, que cobre o ciclo de gestão de uma farmácia desde a entrada de mercadoria "
  "até à emissão do recibo e à produção de relatórios. Todos os objectivos específicos "
  "foram cumpridos, e os requisitos funcionais e não funcionais enunciados no capítulo 3 "
  "encontram-se implementados e verificados.")
p("Do trabalho realizado, três conclusões merecem destaque. A primeira é que a "
  "correcção do abate de stock não é uma questão de algoritmo, mas de garantias "
  "transaccionais. Ordenar os lotes por validade é trivial; o que é difícil, e o que "
  "distingue um sistema fiável de um sistema aparentemente correcto, é assegurar que "
  "essa leitura permanece válida até ao momento da escrita. A combinação de transacção "
  "com bloqueio explícito de linhas é o que torna o sistema seguro em utilização "
  "concorrente.")
p("A segunda é que a integridade deve ser imposta na camada mais profunda possível. As "
  "restrições de verificação declaradas no esquema funcionam como uma última linha de "
  "defesa que nenhum defeito da aplicação consegue contornar.")
p("A terceira é que a decisão de implementar a arquitectura MVC de raiz, ainda que mais "
  "trabalhosa, revelou-se pedagogicamente compensadora: obrigou a resolver "
  "explicitamente problemas que uma framework resolveria de forma opaca, entre eles o "
  "encaminhamento, a protecção contra atribuição em massa e a gestão do ciclo de vida da "
  "sessão.")

doc.add_heading("7.2. Limitações", level=2)
lista([
    "A aplicação não integra facturação electrónica nem comunicação com a autoridade "
    "tributária, pelo que o recibo emitido tem valor meramente interno;",
    "Não existe suporte para múltiplas filiais nem transferência de stock entre "
    "estabelecimentos;",
    "O registo de auditoria cobre os movimentos de stock, mas não as alterações "
    "efectuadas ao catálogo;",
    "A aplicação não dispõe de interface para dispositivos de leitura de código de "
    "barras.",
])

doc.add_heading("7.3. Trabalho futuro", level=2)
lista([
    "Leitura de código de barras no registo de venda e na recepção de mercadoria, "
    "reduzindo o tempo de atendimento e os erros de digitação;",
    "Módulo de encomendas a fornecedores, com sugestão automática de reposição a partir "
    "do stock mínimo e do histórico de consumo;",
    "Alertas automáticos por correio electrónico para lotes próximos da validade;",
    "Exportação dos relatórios em PDF e em folha de cálculo;",
    "Registo de auditoria alargado a todas as alterações de catálogo e de configuração;",
    "Interface de programação para integração com sistemas de facturação e com "
    "plataformas de pagamento móvel.",
])

# ==================================================== 8. REFERÊNCIAS ========

doc.add_heading("8. Referências bibliográficas", level=1)
referencias = [
    "ELMASRI, R.; NAVATHE, S. B. Fundamentals of Database Systems. 7.ª ed. Boston: "
    "Pearson, 2016.",
    "FOWLER, M. Patterns of Enterprise Application Architecture. Boston: "
    "Addison-Wesley, 2003.",
    "GAMMA, E. et al. Design Patterns: Elements of Reusable Object-Oriented Software. "
    "Reading: Addison-Wesley, 1994.",
    "OWASP FOUNDATION. OWASP Top 10 - The Ten Most Critical Web Application Security "
    "Risks. 2021. Disponível em: https://owasp.org/Top10/",
    "OWASP FOUNDATION. Cheat Sheet Series: SQL Injection Prevention, Cross-Site Request "
    "Forgery Prevention, Password Storage. Disponível em: https://cheatsheetseries.owasp.org/",
    "PHP GROUP. PHP Manual: PDO, Sessions, Password Hashing. Disponível em: "
    "https://www.php.net/manual/",
    "ORACLE CORPORATION. MySQL 8.0 Reference Manual: InnoDB Locking and Transaction "
    "Model. Disponível em: https://dev.mysql.com/doc/",
    "OBJECT MANAGEMENT GROUP. OMG Unified Modeling Language (OMG UML), versão 2.5.1. "
    "2017. Disponível em: https://www.omg.org/spec/UML/",
    "BOOTSTRAP TEAM. Bootstrap 5.3 Documentation. Disponível em: "
    "https://getbootstrap.com/docs/5.3/",
    "ORGANIZAÇÃO MUNDIAL DA SAÚDE. Guidelines for the Storage of Essential Medicines and "
    "Other Health Commodities. Genebra: OMS, 2003.",
]
for ref in sorted(referencias):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(1.0)
    par.paragraph_format.first_line_indent = Cm(-1.0)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.line_spacing = 1.15
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.add_run(ref)

# ======================================================== ANEXOS ============

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading("Anexos", level=1)

doc.add_heading("Anexo A - Instalação e execução", level=2)
p("Requisitos: PHP 8.1 ou superior com as extensões pdo_mysql e mbstring activas, e "
  "MySQL 5.7 ou MariaDB 10.3 ou superior. A instalação de referência foi feita sobre "
  "XAMPP.")
codigo(
    "1. Criar a base de dados e as tabelas\n"
    "   mysql -u root < database/schema.sql\n"
    "\n"
    "2. Carregar os dados iniciais (utilizadores, catalogo, lotes)\n"
    "   mysql -u root farmagest < database/seed.sql\n"
    "\n"
    "3. (Opcional) Gerar historico de vendas de demonstracao\n"
    "   php database/seed_vendas.php\n"
    "\n"
    "4. Configurar o acesso a base de dados\n"
    "   copiar .env.example para .env e ajustar DB_HOST, DB_PORT,\n"
    "   DB_NAME, DB_USER e DB_PASS\n"
    "\n"
    "5. Arrancar a aplicacao\n"
    "   php -S localhost:8000 -t public public/router.php\n"
    "   e abrir http://localhost:8000"
)
p("Em alternativa, com o Apache do XAMPP, basta colocar a pasta do projecto dentro de "
  "htdocs e abrir o endereço correspondente. O ficheiro .htaccess existente na raiz do "
  "projecto encaminha automaticamente os pedidos para a pasta public, pelo que não é "
  "necessário incluir /public no endereço:")
codigo(
    "C:\\xampp\\htdocs\\farmagest\\          ->  http://localhost/farmagest/\n"
    "\n"
    "As tres formas de instalacao sao suportadas em simultaneo:\n"
    "   http://localhost:8000/               servidor embutido do PHP\n"
    "   http://localhost/farmagest/          XAMPP, enderecos limpos\n"
    "   http://localhost/farmagest/public/   XAMPP, sem reescrita de URLs"
)
p("Em produção, a raiz do servidor Web deve apontar directamente para a pasta public, "
  "ficando todo o restante código fora da árvore acessível por HTTP. Deve ainda "
  "definir-se APP_ENV=producao e APP_DEBUG=0 no ficheiro de ambiente.")

doc.add_heading("Anexo B - Contas de demonstração", level=2)
p("As contas seguintes são criadas pelo povoamento inicial e destinam-se exclusivamente "
  "a demonstração. Devem ser eliminadas ou ter as palavras-passe alteradas antes de "
  "qualquer utilização real.")
tabela(
    ["Perfil", "Endereço de correio electrónico", "Palavra-passe"],
    [
        ["Administrador", "admin@farmagest.co.mz", "Admin@123"],
        ["Farmacêutico", "farmaceutico@farmagest.co.mz", "Farm@123"],
        ["Operador de caixa", "caixa@farmagest.co.mz", "Caixa@123"],
    ],
    "Contas de demonstração criadas pelo povoamento inicial",
    larguras=[4.0, 7.0, 4.0],
)

doc.add_heading("Anexo C - Verificação da instalação", level=2)
p("O projecto inclui dois procedimentos de verificação que podem ser executados após a "
  "instalação para confirmar que o sistema está operacional:")
codigo(
    "php tests\\executar.php\n"
    "   48 assercoes sobre a logica de negocio\n"
    "\n"
    "powershell -File tests\\crawl.ps1\n"
    "   50 rotas verificadas por HTTP no servidor embutido do PHP\n"
    "\n"
    "powershell -File tests\\crawl.ps1 -Base http://localhost/farmagest\n"
    "   as mesmas 50 rotas verificadas na instalacao do XAMPP"
)
p("Ambos devem terminar sem qualquer falha. O primeiro procedimento repõe o estado da "
  "base de dados no final, pelo que pode ser executado repetidamente sem efeitos "
  "colaterais. O segundo aceita o endereço de base como parâmetro, permitindo verificar "
  "qualquer uma das formas de instalação.")

# ------------------------------------------------------------------ gravar --

os.makedirs(os.path.dirname(SAIDA), exist_ok=True)
doc.save(SAIDA)
print("Relatorio gravado em: %s" % SAIDA)
print("Figuras: %d | Tabelas: %d" % (figura["n"], tabela_n["n"]))
